"""Génération d'offres DOCX à partir du modèle offre_de_prix_BOSS.docx."""

from __future__ import annotations

import re
from io import BytesIO
from pathlib import Path

from docx import Document

from app.config import settings
from app.models.quote import OffreDetail, OffreDocxPayload, QuoteLineItemDashboard

TEMPLATE_PATH = settings.docx_template_path
OUTPUT_DIR = Path(__file__).resolve().parents[2] / "output" / "offres"

# Table 3 : en-tête R0, lignes R1–R15, totaux R16–R18
LINE_TABLE_INDEX = 3
LINE_FIRST_ROW = 1
LINE_LAST_ROW = 15
TOTAL_HT_ROW = 16
TOTAL_TVA_ROW = 17
TOTAL_TTC_ROW = 18

DEST_TABLE_INDEX = 1
REMARKS_TABLE_INDEX = 4


def _format_eur(amount: float) -> str:
    formatted = f"{amount:,.2f}".replace(",", " ").replace(".", ",")
    return f"{formatted} €"


def _safe_filename(numero_offre: str) -> str:
    cleaned = re.sub(r"[^\w\-]+", "_", numero_offre.strip())
    return cleaned or "offre_BOSS"


def _set_cell_text(cell, text: str) -> None:
    cell.text = text


def _offre_to_docx_payload(offre: OffreDetail | OffreDocxPayload) -> OffreDocxPayload:
    if isinstance(offre, OffreDocxPayload):
        return offre
    return OffreDocxPayload(
        numero_offre=offre.numero_offre,
        expediteur_nom=offre.expediteur_nom,
        expediteur_email=offre.expediteur_email,
        sujet=offre.sujet,
        created_at=offre.created_at,
        texte_intro=offre.texte_intro,
        texte_conclusion=offre.texte_conclusion,
        montant_ht=offre.montant_ht,
        montant_tva=offre.montant_tva,
        montant_ttc=offre.montant_ttc,
        lignes=offre.lignes,
    )


def _format_date(value: datetime | str) -> str:
    if hasattr(value, "strftime"):
        return value.strftime("%d/%m/%Y")
    return str(value)[:10]


def _fill_destinataire(doc: Document, offre: OffreDocxPayload) -> None:
    table = doc.tables[DEST_TABLE_INDEX]
    block = (
        f"Offre n° : {offre.numero_offre}\n"
        f"Société  : {offre.expediteur_nom}\n"
        f"Contact  : {offre.expediteur_email}\n"
        f"Objet    : {offre.sujet}\n"
        f"Date     : {_format_date(offre.created_at)}"
    )
    _set_cell_text(table.rows[1].cells[1], block)


def _fill_line_row(row, line: QuoteLineItemDashboard) -> None:
    cells = row.cells
    prix_unitaire = line.prix_unitaire_ht * (1 - line.remise_pct / 100)
    _set_cell_text(cells[0], line.reference)
    _set_cell_text(cells[1], line.designation)
    _set_cell_text(cells[2], str(line.conditionnement))
    _set_cell_text(cells[3], line.alliage)
    _set_cell_text(cells[4], _format_eur(prix_unitaire))
    _set_cell_text(cells[5], "Oui" if line.en_stock else "Non")
    _set_cell_text(cells[6], line.delai_livraison)
    _set_cell_text(cells[7], _format_eur(line.prix_total_ht))


def _fill_lines_and_totals(doc: Document, offre: OffreDocxPayload) -> None:
    table = doc.tables[LINE_TABLE_INDEX]

    for row_idx in range(LINE_FIRST_ROW, LINE_LAST_ROW + 1):
        for col_idx in range(len(table.rows[row_idx].cells)):
            _set_cell_text(table.rows[row_idx].cells[col_idx], "")

    for i, line in enumerate(offre.lignes):
        row_idx = LINE_FIRST_ROW + i
        if row_idx > LINE_LAST_ROW:
            break
        _fill_line_row(table.rows[row_idx], line)

    _set_cell_text(table.rows[TOTAL_HT_ROW].cells[-1], _format_eur(offre.montant_ht))
    _set_cell_text(table.rows[TOTAL_TVA_ROW].cells[-1], _format_eur(offre.montant_tva))
    _set_cell_text(table.rows[TOTAL_TTC_ROW].cells[-1], _format_eur(offre.montant_ttc))


def _fill_remarks(doc: Document, offre: OffreDocxPayload) -> None:
    table = doc.tables[REMARKS_TABLE_INDEX]
    default = table.rows[1].cells[0].text.strip()
    intro = offre.texte_intro.strip()
    conclusion = offre.texte_conclusion.strip()
    parts = [part for part in [intro, default, conclusion] if part]
    _set_cell_text(table.rows[1].cells[0], "\n\n".join(parts))


def _insert_intro_paragraph(doc: Document, offre: OffreDocxPayload) -> None:
    """Ajoute le numéro d'offre sous le titre principal."""
    if not doc.paragraphs:
        return
    title = doc.paragraphs[0]
    if offre.numero_offre and offre.numero_offre not in title.text:
        title.text = f"OFFRE DE PRIX — {offre.numero_offre}"


def generate_offre_docx(
    offre: OffreDetail | OffreDocxPayload,
    *,
    output_path: Path | None = None,
) -> Path:
    payload = _offre_to_docx_payload(offre)
    """
    Remplit le modèle Word BOSS et enregistre le fichier.
    Retourne le chemin du DOCX généré.
    """
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"Modèle introuvable : {TEMPLATE_PATH}")

    doc = Document(TEMPLATE_PATH)
    _insert_intro_paragraph(doc, payload)
    _fill_destinataire(doc, payload)
    _fill_lines_and_totals(doc, payload)
    _fill_remarks(doc, payload)

    if output_path is None:
        OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
        output_path = OUTPUT_DIR / f"{_safe_filename(payload.numero_offre)}.docx"

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


def generate_offre_docx_bytes(offre: OffreDetail | OffreDocxPayload) -> bytes:
    payload = _offre_to_docx_payload(offre)
    """Génère le DOCX en mémoire (endpoint de téléchargement)."""
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"Modèle introuvable : {TEMPLATE_PATH}")

    doc = Document(TEMPLATE_PATH)
    _insert_intro_paragraph(doc, payload)
    _fill_destinataire(doc, payload)
    _fill_lines_and_totals(doc, payload)
    _fill_remarks(doc, payload)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
